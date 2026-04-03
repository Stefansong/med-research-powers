#!/usr/bin/env python3
"""
MRP Manuscript Export: Markdown → .docx
========================================
Converts manuscript markdown files to journal-formatted .docx for submission.

Usage:
    python export_docx.py --manuscript-dir ./manuscript --journal nature --output manuscript.docx
    python export_docx.py --manuscript-dir ./manuscript --journal european-urology --output manuscript.docx

Requirements:
    pip install python-docx pyyaml
"""

import argparse
import os
import re
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── Journal Format Configurations ──────────────────────────────────────────

# Fallback configs for common journal families (used when yaml not available)
JOURNAL_FAMILIES = {
    "nature": {
        "font": "Times New Roman",
        "font_size": 12,
        "heading1_size": 14,
        "heading2_size": 13,
        "line_spacing": 2.0,
        "margin_cm": 2.54,
        "section_order": [
            "title-page", "abstract", "introduction", "results",
            "discussion", "methods", "references"
        ],
        "special": ["reporting_summary"],
    },
    "lancet": {
        "font": "Times New Roman",
        "font_size": 12,
        "heading1_size": 14,
        "heading2_size": 13,
        "line_spacing": 2.0,
        "margin_cm": 2.54,
        "section_order": [
            "title-page", "abstract", "research-in-context",
            "introduction", "methods", "results", "discussion", "references"
        ],
        "special": ["research_in_context"],
    },
    "jama": {
        "font": "Times New Roman",
        "font_size": 11,
        "heading1_size": 13,
        "heading2_size": 12,
        "line_spacing": 2.0,
        "margin_cm": 2.54,
        "section_order": [
            "title-page", "key-points", "abstract", "introduction",
            "methods", "results", "discussion", "references"
        ],
        "special": ["key_points"],
    },
    "standard": {
        "font": "Times New Roman",
        "font_size": 12,
        "heading1_size": 14,
        "heading2_size": 13,
        "line_spacing": 2.0,
        "margin_cm": 2.54,
        "section_order": [
            "title-page", "abstract", "introduction", "methods",
            "results", "discussion", "references"
        ],
        "special": [],
    },
    "ieee": {
        "font": "Times New Roman",
        "font_size": 10,
        "heading1_size": 12,
        "heading2_size": 11,
        "line_spacing": 1.0,
        "margin_cm": 1.91,
        "section_order": [
            "title-page", "abstract", "introduction", "related-work",
            "methods", "results", "discussion", "references"
        ],
        "special": ["index_terms"],
    },
}

# Map journal IDs to families
JOURNAL_FAMILY_MAP = {
    # Nature family
    "npj-digital-surgery": "nature",
    "npj-digital-medicine": "nature",
    "nature-medicine": "nature",
    "nature-communications": "nature",
    "scientific-reports": "nature",
    # Lancet family
    "lancet": "lancet",
    "lancet-digital-health": "lancet",
    "lancet-oncology": "lancet",
    # JAMA family
    "jama": "jama",
    "jama-surgery": "jama",
    "jama-oncology": "jama",
    # Standard
    "european-urology": "standard",
    "journal-urology": "standard",
    "bju-international": "standard",
    "annals-of-surgery": "standard",
    "surgical-endoscopy": "standard",
    "international-journal-surgery": "standard",
    "journal-endourology": "standard",
    # AI/Tech
    "medical-image-analysis": "standard",
    "jmir": "standard",
    "ieee-jbhi": "ieee",
    "ieee-tmi": "ieee",
}


def get_journal_config(journal_id: str, yaml_path: str = None) -> dict:
    """Load journal configuration from yaml or fallback to built-in."""
    # Try loading from yaml first
    if yaml_path and os.path.exists(yaml_path):
        with open(yaml_path, "r") as f:
            templates = yaml.safe_load(f)
        for entry in templates.get("journals", []):
            if entry.get("id") == journal_id:
                family = JOURNAL_FAMILY_MAP.get(journal_id, "standard")
                config = JOURNAL_FAMILIES[family].copy()
                # Override with yaml specifics
                if entry.get("word_limit"):
                    config["word_limit"] = entry["word_limit"]
                if entry.get("references"):
                    config["ref_limit"] = entry["references"]
                config["journal_name"] = entry.get("journal", journal_id)
                return config

    # Fallback to family mapping
    family = JOURNAL_FAMILY_MAP.get(journal_id, "standard")
    config = JOURNAL_FAMILIES[family].copy()
    config["journal_name"] = journal_id
    return config


# ─── Markdown Parsing ────────────────────────────────────────────────────────

def strip_html_comments(text: str) -> str:
    """Remove HTML comments <!-- ... -->."""
    return re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)


def parse_markdown_line(line: str) -> dict:
    """Parse a single markdown line into type and content."""
    stripped = line.strip()

    if not stripped:
        return {"type": "blank"}

    # Headings
    match = re.match(r'^(#{1,3})\s+(.*)', stripped)
    if match:
        level = len(match.group(1))
        return {"type": f"heading{level}", "text": match.group(2)}

    # Table row
    if stripped.startswith("|") and stripped.endswith("|"):
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        # Check if separator row
        if all(re.match(r'^[-:]+$', c) for c in cells):
            return {"type": "table_separator"}
        return {"type": "table_row", "cells": cells}

    # Bullet list
    if re.match(r'^[-*]\s+', stripped):
        text = re.sub(r'^[-*]\s+', '', stripped)
        return {"type": "bullet", "text": text}

    # Numbered list
    match = re.match(r'^(\d+)\.\s+(.*)', stripped)
    if match:
        return {"type": "numbered", "number": int(match.group(1)), "text": match.group(2)}

    # Block quote
    if stripped.startswith(">"):
        text = stripped.lstrip("> ")
        return {"type": "quote", "text": text}

    # Regular paragraph
    return {"type": "paragraph", "text": stripped}


def add_formatted_text(paragraph, text: str, config: dict, base_bold=False, base_italic=False):
    """Add text with inline markdown formatting (bold, italic, superscript)."""
    # Pattern for **bold**, *italic*, ^superscript^, ~subscript~
    pattern = r'(\*\*.*?\*\*|\*.*?\*|\^.*?\^|~.*?~|[^*^~]+)'
    parts = re.findall(pattern, text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = config["font"]
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = config["font"]
        elif part.startswith("^") and part.endswith("^"):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.name = config["font"]
        elif part.startswith("~") and part.endswith("~"):
            run = paragraph.add_run(part[1:-1])
            run.font.subscript = True
            run.font.name = config["font"]
        else:
            run = paragraph.add_run(part)
            run.font.name = config["font"]
            if base_bold:
                run.bold = True
            if base_italic:
                run.italic = True


def set_cell_shading(cell, color_hex: str):
    """Set cell background shading."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


# ─── Document Builder ────────────────────────────────────────────────────────

def build_docx(manuscript_dir: str, config: dict, output_path: str) -> dict:
    """Build .docx from manuscript markdown files."""
    doc = Document()
    stats = {"word_count": 0, "ref_count": 0, "fig_count": 0, "table_count": 0,
             "pages": 0, "placeholders": [], "missing_sections": [], "sections_included": []}

    # ─── Document Setup ───
    style = doc.styles['Normal']
    style.font.name = config["font"]
    style.font.size = Pt(config["font_size"])
    style.paragraph_format.line_spacing = config["line_spacing"]
    style.paragraph_format.space_after = Pt(0)

    # CJK font fallback
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{config["font"]}"/>')
        rPr.append(rFonts)

    for section in doc.sections:
        section.top_margin = Cm(config["margin_cm"])
        section.bottom_margin = Cm(config["margin_cm"])
        section.left_margin = Cm(config["margin_cm"])
        section.right_margin = Cm(config["margin_cm"])

    # Configure heading styles
    for level, size_key in [(1, "heading1_size"), (2, "heading2_size")]:
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = config["font"]
        h_style.font.size = Pt(config.get(size_key, 14))
        h_style.font.bold = True
        h_style.font.color.rgb = RGBColor(0, 0, 0)

    # ─── Process Sections ───
    section_file_map = {
        "title-page": "title-page.md",
        "abstract": "abstract.md",
        "introduction": "introduction.md",
        "results": "results.md",
        "discussion": "discussion.md",
        "methods": "methods.md",
        "references": "references.md",
        "supplementary": "supplementary.md",
    }

    first_section = True
    for section_id in config["section_order"]:
        filename = section_file_map.get(section_id)
        if not filename:
            continue

        filepath = os.path.join(manuscript_dir, filename)
        if not os.path.exists(filepath):
            stats["missing_sections"].append(section_id)
            continue

        stats["sections_included"].append(section_id)

        # Page break between sections (not before first)
        if not first_section:
            doc.add_page_break()
        first_section = False

        # Read and process markdown
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()

        content = strip_html_comments(content)
        lines = content.split("\n")

        # Track table state
        in_table = False
        table_headers = []
        table_rows = []

        for line in lines:
            parsed = parse_markdown_line(line)

            # Flush table if leaving table context
            if parsed["type"] not in ("table_row", "table_separator") and in_table:
                _add_table(doc, table_headers, table_rows, config)
                stats["table_count"] += 1
                in_table = False
                table_headers = []
                table_rows = []

            if parsed["type"] == "blank":
                continue

            elif parsed["type"] == "heading1":
                doc.add_heading(parsed["text"], level=1)

            elif parsed["type"] == "heading2":
                doc.add_heading(parsed["text"], level=2)

            elif parsed["type"] == "heading3":
                doc.add_heading(parsed["text"], level=3)

            elif parsed["type"] == "paragraph":
                p = doc.add_paragraph()
                add_formatted_text(p, parsed["text"], config)
                stats["word_count"] += len(parsed["text"].split())
                # Check for placeholders
                for marker in ["[pending]", "[TBD]", "[TODO]", "PLACEHOLDER", "[N]"]:
                    if marker.lower() in parsed["text"].lower():
                        stats["placeholders"].append(f"{filename}: {parsed['text'][:80]}")

            elif parsed["type"] == "bullet":
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, parsed["text"], config)
                stats["word_count"] += len(parsed["text"].split())

            elif parsed["type"] == "numbered":
                p = doc.add_paragraph(style='List Number')
                add_formatted_text(p, parsed["text"], config)
                stats["word_count"] += len(parsed["text"].split())

            elif parsed["type"] == "quote":
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.27)
                add_formatted_text(p, parsed["text"], config, base_italic=True)

            elif parsed["type"] == "table_row":
                if not in_table:
                    table_headers = parsed["cells"]
                    in_table = True
                else:
                    table_rows.append(parsed["cells"])

            elif parsed["type"] == "table_separator":
                pass  # Skip separator rows

        # Flush any remaining table
        if in_table:
            _add_table(doc, table_headers, table_rows, config)
            stats["table_count"] += 1

    # Count references
    ref_path = os.path.join(manuscript_dir, "references.md")
    if os.path.exists(ref_path):
        with open(ref_path, "r") as f:
            ref_content = f.read()
        stats["ref_count"] = len(re.findall(r'^\d+\.', ref_content, re.MULTILINE))

    # Save
    doc.save(output_path)
    return stats


def _add_table(doc, headers, rows, config):
    """Add a formatted table to the document."""
    if not headers:
        return

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = config["font"]
        run.font.size = Pt(config["font_size"] - 1)
        set_cell_shading(cell, "D9D9D9")

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx in range(min(len(row_data), n_cols)):
            cell = table.rows[row_idx + 1].cells[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(row_data[col_idx])
            run.font.name = config["font"]
            run.font.size = Pt(config["font_size"] - 1)

    doc.add_paragraph()  # spacing after table


# ─── Report Generation ───────────────────────────────────────────────────────

def generate_report(stats: dict, config: dict, output_path: str) -> str:
    """Generate export quality report."""
    word_limit = config.get("word_limit")
    ref_limit = config.get("ref_limit")

    word_status = ""
    if word_limit:
        if isinstance(word_limit, str) and "no" in word_limit.lower():
            word_status = "✅ No limit"
        else:
            try:
                limit = int(str(word_limit).replace(",", ""))
                word_status = f"✅ {stats['word_count']}/{limit}" if stats['word_count'] <= limit \
                    else f"⚠️ {stats['word_count']}/{limit} — 超出 {stats['word_count']-limit} 字"
            except ValueError:
                word_status = f"ℹ️ {stats['word_count']} (limit: {word_limit})"
    else:
        word_status = f"ℹ️ {stats['word_count']} (no limit specified)"

    ref_status = ""
    if ref_limit:
        try:
            limit = int(str(ref_limit).split(",")[0].replace("≤", "").strip())
            ref_status = f"✅ {stats['ref_count']}/{limit}" if stats['ref_count'] <= limit \
                else f"⚠️ {stats['ref_count']}/{limit}"
        except ValueError:
            ref_status = f"ℹ️ {stats['ref_count']}"
    else:
        ref_status = f"ℹ️ {stats['ref_count']}"

    report = f"""# Manuscript Export Report

## File
- **Output:** `{output_path}`
- **Journal:** {config.get('journal_name', 'Unknown')}
- **Format family:** {config.get('font')} {config.get('font_size')}pt, {config.get('line_spacing')}x spacing

## Statistics
- **Word count:** {word_status}
- **References:** {ref_status}
- **Tables:** {stats['table_count']}
- **Sections included:** {', '.join(stats['sections_included'])}
- **Sections missing:** {', '.join(stats['missing_sections']) if stats['missing_sections'] else 'None'}

## Placeholder Warnings
"""
    if stats["placeholders"]:
        for ph in stats["placeholders"][:20]:
            report += f"- ⚠️ {ph}\n"
    else:
        report += "- ✅ No placeholders detected\n"

    return report


# ─── Entry Point ─────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="MRP Manuscript Export")
    parser.add_argument("--manuscript-dir", required=True, help="Directory containing .md files")
    parser.add_argument("--journal", default="standard", help="Journal ID or family name")
    parser.add_argument("--output", default="manuscript.docx", help="Output .docx path")
    parser.add_argument("--yaml", default=None, help="Path to journal-templates.yaml")
    parser.add_argument("--report", default=None, help="Path for export report .md")
    args = parser.parse_args()

    config = get_journal_config(args.journal, args.yaml)
    print(f"Exporting for: {config.get('journal_name', args.journal)}")
    print(f"Format: {config['font']} {config['font_size']}pt, {config['line_spacing']}x spacing")
    print(f"Section order: {' → '.join(config['section_order'])}")

    stats = build_docx(args.manuscript_dir, config, args.output)

    report = generate_report(stats, config, args.output)
    print(report)

    if args.report:
        with open(args.report, "w") as f:
            f.write(report)
        print(f"Report saved to {args.report}")


if __name__ == "__main__":
    main()
